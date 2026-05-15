import os
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles parsing and extraction from various document formats"""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> List[Dict[str, any]]:
        """
        Extract text from PDF file page by page
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of dicts with page number and text content
        """
        chunks = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():  # Only add non-empty pages
                        chunks.append({
                            "page": page_num + 1,
                            "text": text.strip(),
                            "metadata": {
                                "source": os.path.basename(file_path),
                                "page": page_num + 1,
                                "total_pages": total_pages
                            }
                        })
                        
            logger.info(f"Extracted {len(chunks)} pages from PDF: {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {str(e)}")
            raise

    @staticmethod
    def extract_text_from_docx(file_path: str) -> List[Dict[str, any]]:
        """
        Extract text from DOCX file paragraph by paragraph
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of dicts with paragraph number and text content
        """
        chunks = []
        try:
            doc = Document(file_path)
            
            for idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                if text:  # Only add non-empty paragraphs
                    chunks.append({
                        "page": idx + 1,  # Using paragraph number as "page"
                        "text": text,
                        "metadata": {
                            "source": os.path.basename(file_path),
                            "paragraph": idx + 1,
                            "total_paragraphs": len(doc.paragraphs)
                        }
                    })
                    
            logger.info(f"Extracted {len(chunks)} paragraphs from DOCX: {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {str(e)}")
            raise

    @staticmethod
    def extract_text_from_txt(file_path: str) -> List[Dict[str, any]]:
        """
        Extract text from TXT file, chunking by lines
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            List of dicts with line chunks
        """
        chunks = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                # Split into chunks of ~500 characters for better retrieval
                chunk_size = 500
                lines = content.split('\n')
                
                current_chunk = ""
                chunk_num = 1
                
                for line in lines:
                    if len(current_chunk) + len(line) < chunk_size:
                        current_chunk += line + "\n"
                    else:
                        if current_chunk.strip():
                            chunks.append({
                                "page": chunk_num,
                                "text": current_chunk.strip(),
                                "metadata": {
                                    "source": os.path.basename(file_path),
                                    "chunk": chunk_num
                                }
                            })
                        current_chunk = line + "\n"
                        chunk_num += 1
                
                # Add final chunk
                if current_chunk.strip():
                    chunks.append({
                        "page": chunk_num,
                        "text": current_chunk.strip(),
                        "metadata": {
                            "source": os.path.basename(file_path),
                            "chunk": chunk_num
                        }
                    })
                    
            logger.info(f"Extracted {len(chunks)} chunks from TXT: {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error extracting TXT {file_path}: {str(e)}")
            raise

    @staticmethod
    def process_document(file_path: str) -> List[Dict[str, any]]:
        """
        Main processing method - detects format and extracts text
        
        Args:
            file_path: Path to document
            
        Returns:
            List of text chunks with metadata
        """
        extension = Path(file_path).suffix.lower()
        
        if extension == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif extension == '.txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
